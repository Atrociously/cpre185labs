import sys
import os
from docx import Document
from docx.shared import Inches

def main():
    if len(sys.argv) < 3:
        print("Need an output file and an input directory")
        return

    outfile = sys.argv[1]
    indir = sys.argv[2]

    doc = Document()

    for file in os.scandir(indir):
        if not file.name.endswith(".c"):
            continue
        with open(file, 'r') as f:
            code = f.read()
            doc.add_paragraph(code)
        doc.add_page_break()

    for file in os.scandir(indir):
        if not (file.name.endswith(".jpg") or file.name.endswith(".png")):
            continue
        doc.add_picture(file.path, width=Inches(6))


    doc.save(outfile)


if __name__ == "__main__":
    main()
